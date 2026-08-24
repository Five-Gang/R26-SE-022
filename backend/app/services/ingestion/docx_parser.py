from __future__ import annotations
"""Docx Parser — extracts text and metadata from Word (.docx) documents."""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Union

import docx


@dataclass
class DocxExtractionResult:
    """Result of parsing a .docx file."""

    full_text: str
    paragraphs: list[str] = field(default_factory=list)
    tables_text: list[str] = field(default_factory=list)
    headings: list[dict] = field(default_factory=list)


class DocxParser:
    """Extracts structured text from Word (.docx) files."""

    def extract(self, file_path: Union[str, Path]) -> DocxExtractionResult:
        """Extract text, headings, and tables from a .docx file.

        Args:
            file_path: Path to the .docx file.

        Returns:
            DocxExtractionResult containing full_text, paragraphs, and headings.
        """
        doc = docx.Document(str(file_path))

        paragraphs = []
        headings = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            paragraphs.append(text)

            if p.style.name.startswith("Heading"):
                level = 1
                try:
                    level = int(p.style.name.replace("Heading", "").strip())
                except ValueError:
                    pass
                headings.append({"title": text, "level": level})

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                tables_text.append("\n".join(table_rows))

        all_content = []
        if paragraphs:
            all_content.extend(paragraphs)
        if tables_text:
            all_content.extend(tables_text)

        full_text = "\n\n".join(all_content)

        return DocxExtractionResult(
            full_text=full_text,
            paragraphs=paragraphs,
            tables_text=tables_text,
            headings=headings,
        )
